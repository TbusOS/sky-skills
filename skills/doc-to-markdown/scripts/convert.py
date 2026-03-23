#!/usr/bin/env python3
"""
Convert PDF and DOCX files to Markdown with extracted images.

Usage:
    python3 convert.py <source>                     # single file or directory
    python3 convert.py <source> --output <out-dir>  # specify output directory

Output structure:
    <out-dir>/
    ├── <doc-name>.md
    └── images/
        └── <doc-slug>/
            ├── sec01_01_overview.png
            └── sec02_01_boot_flow.png
"""

import sys
import os
import re
import zipfile
import argparse


def slugify(name):
    """Create a filesystem-safe slug from a filename."""
    name = os.path.splitext(name)[0]
    name = re.sub(r'[^\w\u4e00-\u9fff\u3400-\u4dbf-]', '_', name)
    name = re.sub(r'_+', '_', name).strip('_')
    return name


def make_img_name(section_num, img_seq, desc=''):
    """Generate image filename: sec{NN}_{seq}_{desc}.png"""
    sec = f"sec{section_num:02d}" if section_num else "sec00"
    seq = f"{img_seq:02d}"
    if desc:
        desc = slugify(desc)[:40]
        return f"{sec}_{seq}_{desc}.png"
    return f"{sec}_{seq}.png"


def ensure_deps():
    """Check required dependencies."""
    missing = []
    try:
        import fitz  # noqa: F401
    except ImportError:
        missing.append('pymupdf')
    try:
        from docx import Document  # noqa: F401
    except ImportError:
        missing.append('python-docx')
    if missing:
        print(f"Missing: {', '.join(missing)}")
        print(f"Install: pip3 install {' '.join(missing)}")
        sys.exit(1)


def try_convert_emf(emf_data, out_path):
    """Try to convert EMF/WMF to PNG. Returns True on success."""
    import shutil
    import subprocess
    import tempfile

    if shutil.which('libreoffice'):
        try:
            with tempfile.NamedTemporaryFile(suffix='.emf', delete=False) as tmp:
                tmp.write(emf_data)
                tmp_path = tmp.name
            out_dir = os.path.dirname(out_path)
            subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'png', tmp_path, '--outdir', out_dir],
                capture_output=True, timeout=30
            )
            converted = os.path.join(out_dir, os.path.splitext(os.path.basename(tmp_path))[0] + '.png')
            if os.path.exists(converted):
                os.rename(converted, out_path)
                os.unlink(tmp_path)
                return True
            os.unlink(tmp_path)
        except Exception:
            pass

    try:
        from PIL import Image
        import io
        img = Image.open(io.BytesIO(emf_data))
        img.save(out_path, 'PNG')
        return True
    except Exception:
        pass

    return False


# Minimum image size to keep (skip tiny icons/decorative images)
MIN_IMG_SIZE = 5000  # bytes


def detect_section_from_text(text):
    """Scan all lines in text to find the top-level section number.
    Returns (section_number, matched_or_not)."""
    if not text:
        return None, False
    for line in text.split('\n'):
        line = line.strip()
        m = re.match(r'^(\d+)(?:\.\d+)*(?:\s|\.)', line)
        if m:
            try:
                return int(m.group(1)), True
            except ValueError:
                pass
    return None, False


def extract_pdf_tables(page):
    """Extract tables from a PDF page using PyMuPDF's find_tables (if available).
    Returns list of markdown-formatted table strings."""
    tables_md = []
    try:
        tabs = page.find_tables()
        for tab in tabs:
            data = tab.extract()
            if not data or len(data) < 1:
                continue
            # Build markdown table
            lines = []
            for i, row in enumerate(data):
                cells = [str(c).replace('\n', ' ').strip() if c else '' for c in row]
                lines.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
            tables_md.append('\n'.join(lines))
    except (AttributeError, Exception):
        # find_tables not available in older PyMuPDF versions
        pass
    return tables_md


def convert_pdf(pdf_path, md_dir, img_dir):
    """Convert a PDF to markdown with images."""
    import fitz

    basename = os.path.basename(pdf_path)
    slug = slugify(basename)
    img_subdir = os.path.join(img_dir, slug)
    os.makedirs(img_subdir, exist_ok=True)

    doc = fitz.open(pdf_path)
    md_lines = [f"# {os.path.splitext(basename)[0]}\n"]
    md_lines.append(f"> Source: `{basename}` ({doc.page_count} pages)\n")
    md_lines.append("---\n")

    image_files = []
    current_section = 0
    img_seq = 0

    for i in range(doc.page_count):
        page = doc[i]
        page_num = i + 1
        text = page.get_text().strip()

        # Detect section number from any line on the page
        sec_num, found = detect_section_from_text(text)
        if found and sec_num != current_section:
            current_section = sec_num
            img_seq = 0

        # Extract embedded images from the page
        img_list = page.get_images(full=True)
        has_significant_image = False

        for img_idx, img_info in enumerate(img_list):
            xref = img_info[0]
            try:
                base_image = doc.extract_image(xref)
                if not base_image:
                    continue
                img_data = base_image["image"]
                img_ext = base_image.get("ext", "png")
                width = base_image.get("width", 0)
                height = base_image.get("height", 0)

                # Skip tiny images (icons, bullets, decorative)
                if len(img_data) < MIN_IMG_SIZE and width < 50 and height < 50:
                    continue

                img_seq += 1
                img_name = make_img_name(current_section, img_seq)
                img_path = os.path.join(img_subdir, img_name)

                # Convert to PNG if not already
                if img_ext.lower() != 'png':
                    try:
                        from PIL import Image
                        import io
                        pil_img = Image.open(io.BytesIO(img_data))
                        pil_img.save(img_path, 'PNG')
                    except Exception:
                        with open(img_path, 'wb') as f:
                            f.write(img_data)
                else:
                    with open(img_path, 'wb') as f:
                        f.write(img_data)

                # Check if image is too small, export a high-res page version too
                if width < 400 or height < 300:
                    # Also export full page as high-res backup
                    img_seq += 1
                    page_img_name = make_img_name(current_section, img_seq, f"page_{page_num}_hires")
                    page_img_path = os.path.join(img_subdir, page_img_name)
                    mat = fitz.Matrix(2.5, 2.5)
                    pix = page.get_pixmap(matrix=mat)
                    pix.save(page_img_path)
                    image_files.append(img_name)
                    image_files.append(page_img_name)

                    rel_img = f"images/{slug}/{img_name}"
                    rel_page_img = f"images/{slug}/{page_img_name}"
                    md_lines.append(f"\n![{img_name}]({rel_img})\n")
                    md_lines.append(f"<!-- High-res page backup: ![Page {page_num}]({rel_page_img}) -->\n")
                    has_significant_image = True
                else:
                    image_files.append(img_name)
                    rel_img = f"images/{slug}/{img_name}"
                    md_lines.append(f"\n![{img_name}]({rel_img})\n")
                    has_significant_image = True

            except Exception:
                pass

        # Extract tables from this page
        tables = extract_pdf_tables(page)

        # Add text content
        if text:
            # Collect table cell texts to skip them from body text
            table_texts = set()
            for tmd in tables:
                for cell in re.findall(r'\|\s*([^|]+?)\s*\|', tmd):
                    cell = cell.strip()
                    if cell and cell != '---':
                        table_texts.add(cell)

            lines = text.split('\n')
            in_table_region = False
            for line in lines:
                line = line.strip()
                if not line:
                    md_lines.append("")
                    continue

                # Skip lines that are table cell content (will be rendered as table below)
                if table_texts and line in table_texts:
                    in_table_region = True
                    continue
                if in_table_region and len(line) < 30:
                    # Still inside fragmented table text
                    if line in table_texts:
                        continue
                    # Check if this could be a straggler
                    for tt in table_texts:
                        if line in tt:
                            continue

                in_table_region = False

                # Detect heading patterns
                heading_match = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)$', line)
                if heading_match:
                    parts = heading_match.group(1).split('.')
                    level = min(len(parts) + 1, 6)
                    md_lines.append(f"{'#' * level} {line}\n")
                else:
                    md_lines.append(line)

            # Insert tables in markdown format
            if tables:
                md_lines.append("")
                for tmd in tables:
                    md_lines.append(tmd)
                    md_lines.append("")

            md_lines.append("")

        # If no embedded images found at all, export full page as image
        if not img_list:
            mat = fitz.Matrix(2, 2)
            pix = page.get_pixmap(matrix=mat)
            img_seq += 1
            img_name = make_img_name(current_section, img_seq, f"page_{page_num}")
            img_path = os.path.join(img_subdir, img_name)
            pix.save(img_path)
            image_files.append(img_name)
            rel_img = f"images/{slug}/{img_name}"
            md_lines.append(f"\n![Page {page_num}]({rel_img})\n")

    page_count = doc.page_count
    doc.close()

    md_path = os.path.join(md_dir, f"{slug}.md")
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(md_lines))

    print(f"  PDF -> {os.path.basename(md_path)} ({page_count} pages, {len(image_files)} images)")
    return md_path, image_files


def convert_docx(docx_path, md_dir, img_dir):
    """Convert a DOCX to markdown with extracted images."""
    from docx import Document

    basename = os.path.basename(docx_path)
    slug = slugify(basename)
    img_subdir = os.path.join(img_dir, slug)
    os.makedirs(img_subdir, exist_ok=True)

    doc = Document(docx_path)

    # --- Phase 1: Extract all images from the docx zip ---
    media_map = {}  # original_fname -> output_img_name
    emf_skipped = []
    emf_converted = []
    img_counter = 0

    zf = zipfile.ZipFile(docx_path)
    rels = {}
    # Parse relationships to map rId -> media file
    for name in zf.namelist():
        if name == 'word/_rels/document.xml.rels':
            import xml.etree.ElementTree as ET
            rels_xml = zf.read(name).decode('utf-8')
            root = ET.fromstring(rels_xml)
            for rel in root:
                rid = rel.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}Id',
                             rel.get('Id', ''))
                target = rel.get('Target', '')
                if target.startswith('media/'):
                    rels[rid] = target

    # Extract media files
    for name in zf.namelist():
        if name.startswith('word/media/'):
            fname = os.path.basename(name)
            ext = os.path.splitext(fname)[1].lower()
            data = zf.read(name)

            img_counter += 1

            if ext in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.svg'):
                img_name = make_img_name(0, img_counter)
                out_path = os.path.join(img_subdir, img_name)
                if ext not in ('.png', '.svg'):
                    try:
                        from PIL import Image
                        import io
                        pil_img = Image.open(io.BytesIO(data))
                        pil_img.save(out_path, 'PNG')
                    except Exception:
                        with open(out_path, 'wb') as f:
                            f.write(data)
                else:
                    with open(out_path, 'wb') as f:
                        f.write(data)
                media_map[fname] = img_name

            elif ext in ('.emf', '.wmf'):
                img_name = make_img_name(0, img_counter)
                out_path = os.path.join(img_subdir, img_name)
                if try_convert_emf(data, out_path):
                    media_map[fname] = img_name
                    emf_converted.append(fname)
                else:
                    emf_skipped.append(fname)
    zf.close()

    # --- Phase 2: Convert document content ---
    md_lines = [f"# {os.path.splitext(basename)[0]}\n"]
    md_lines.append(f"> Source: `{basename}`\n")
    md_lines.append("---\n")

    inserted_images = set()

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''
        text = para.text.strip()

        # Check for embedded images in this paragraph (via XML)
        try:
            import xml.etree.ElementTree as ET
            para_xml = para._element.xml
            if 'blip' in para_xml or 'imagedata' in para_xml:
                root = ET.fromstring(f'<root xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
                                    f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
                                    f'xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture" '
                                    f'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                                    f'xmlns:v="urn:schemas-microsoft-com:vml" '
                                    f'xmlns:o="urn:schemas-microsoft-com:office:office">'
                                    f'{para_xml}</root>')
                for blip in root.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                    embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed', '')
                    if embed and embed in rels:
                        media_file = os.path.basename(rels[embed])
                        if media_file in media_map and media_file not in inserted_images:
                            img_name = media_map[media_file]
                            rel_img = f"images/{slug}/{img_name}"
                            md_lines.append(f"\n![{img_name}]({rel_img})\n")
                            inserted_images.add(media_file)
        except Exception:
            pass

        if not text:
            md_lines.append("")
            continue

        # Format based on style
        if style_name.startswith('Heading'):
            try:
                level = int(style_name.replace('Heading', '').strip())
                level = min(level + 1, 6)
            except ValueError:
                level = 2
            md_lines.append(f"{'#' * level} {text}\n")
        elif style_name == 'Title':
            md_lines.append(f"## {text}\n")
        elif 'List' in style_name or 'Bullet' in style_name:
            md_lines.append(f"- {text}")
        elif 'Number' in style_name:
            md_lines.append(f"1. {text}")
        elif 'Code' in style_name or 'code' in style_name:
            md_lines.append(f"```\n{text}\n```")
        else:
            md_lines.append(text)

    # Convert tables (inline with position detection)
    for table in doc.tables:
        md_lines.append("")
        for i, row in enumerate(table.rows):
            cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            line = "| " + " | ".join(cells) + " |"
            md_lines.append(line)
            if i == 0:
                separator = "| " + " | ".join(["---"] * len(cells)) + " |"
                md_lines.append(separator)
        md_lines.append("")

    # Add any images not yet inserted
    all_inserted_names = {media_map[k] for k in inserted_images if k in media_map}
    remaining_imgs = [(k, v) for k, v in media_map.items() if v not in all_inserted_names]
    if remaining_imgs:
        md_lines.append("\n## Other Images\n")
        for orig, img_name in remaining_imgs:
            rel_img = f"images/{slug}/{img_name}"
            md_lines.append(f"![{img_name}]({rel_img})\n")

    if emf_skipped:
        md_lines.append("\n> **Note**: The following EMF/WMF images could not be converted "
                       "(install LibreOffice for EMF support):")
        for emf in emf_skipped:
            md_lines.append(f"> - `{emf}`")
        md_lines.append("")

    md_path = os.path.join(md_dir, f"{slug}.md")
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(md_lines))

    total_imgs = len(media_map)
    info = f"{total_imgs} images extracted"
    if emf_converted:
        info += f", {len(emf_converted)} EMF converted"
    if emf_skipped:
        info += f", {len(emf_skipped)} EMF skipped"
    print(f"  DOCX -> {os.path.basename(md_path)} ({info})")
    return md_path, list(media_map.values())


def main():
    parser = argparse.ArgumentParser(description='Convert PDF/DOCX to Markdown')
    parser.add_argument('source', help='Source file or directory')
    parser.add_argument('--output', '-o', help='Output directory (default: <project>/markdown/)')
    args = parser.parse_args()

    src = args.source
    ensure_deps()

    # Determine if source is a single file or directory
    if os.path.isfile(src):
        files_to_convert = [src]
        src_dir = os.path.dirname(src) or '.'
    elif os.path.isdir(src):
        src_dir = src
        all_files = os.listdir(src)
        files_to_convert = [
            os.path.join(src, f) for f in sorted(all_files)
            if (f.lower().endswith('.pdf') or f.lower().endswith('.docx'))
            and not f.startswith('~$')
        ]
    else:
        print(f"Error: {src} not found")
        sys.exit(1)

    if not files_to_convert:
        print("No PDF or DOCX files found.")
        sys.exit(0)

    # Output directory
    md_dir = args.output if args.output else os.path.join(src_dir, 'markdown')
    img_dir = os.path.join(md_dir, 'images')
    os.makedirs(img_dir, exist_ok=True)

    pdfs = [f for f in files_to_convert if f.lower().endswith('.pdf')]
    docxs = [f for f in files_to_convert if f.lower().endswith('.docx')]
    print(f"Found {len(pdfs)} PDF(s) and {len(docxs)} DOCX file(s)\n")
    print("Converting...\n")

    results = []
    for fpath in files_to_convert:
        try:
            if fpath.lower().endswith('.pdf'):
                md_path, imgs = convert_pdf(fpath, md_dir, img_dir)
            else:
                md_path, imgs = convert_docx(fpath, md_dir, img_dir)
            results.append((os.path.basename(fpath), md_path, len(imgs), 'ok'))
        except Exception as e:
            print(f"  Error converting {os.path.basename(fpath)}: {e}")
            results.append((os.path.basename(fpath), None, 0, str(e)))

    print(f"\nDone! Output: {md_dir}/")
    ok = sum(1 for r in results if r[3] == 'ok')
    print(f"  {ok}/{len(results)} converted successfully")
    total_imgs = sum(r[2] for r in results)
    print(f"  {total_imgs} total images extracted")

    return 0


if __name__ == '__main__':
    sys.exit(main() or 0)
